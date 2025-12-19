"""
Document Processor for Dropbox Files
Extracts text from PDFs and other documents
Prepares content for indexing in Weaviate
"""

import logging
import io
from typing import Dict, Optional, List, Any
from pathlib import Path
import PyPDF2
import re
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    from docx import Document
    docx_available = True
except ImportError:
    docx_available = False
    logger.warning("python-docx not installed - DOCX support disabled")
try:
    import openpyxl
except ImportError:
    openpyxl = None
    logger.warning("openpyxl not installed - XLSX support disabled")


class DocumentProcessor:
    """
    Processes documents for indexing
    Extracts text, metadata, and structures content
    """
    
    def __init__(self):
        """Initialize document processor"""
        self.supported_extensions = {'.pdf', '.txt', '.md', '.csv'}
        self.processed_count = 0
        self.failed_count = 0
        
    def process_document(self, content: bytes, file_metadata: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """
        Process a document for indexing
        
        Args:
            content: Raw file content
            file_metadata: Metadata from Dropbox
            
        Returns:
            Processed document ready for indexing
        """
        try:
            file_path = file_metadata.get('path_display', '')
            file_ext = Path(file_path).suffix.lower()
            
            # Extract text based on file type
            text = None
            if file_ext == '.pdf':
                text = self._extract_pdf_text(content)
            elif file_ext in ['.txt', '.md', '.csv']:
                text = content.decode('utf-8', errors='ignore')
            elif file_ext in ['.docx', '.doc']:
                text = self._extract_docx_text(content)
            elif file_ext in ['.xlsx', '.xls']:
                text = self._extract_xlsx_text(content)
            else:
                logger.debug(f"Unsupported file type: {file_ext}")
                return None
            
            if not text:
                logger.warning(f"No text extracted from {file_path}")
                return None
            
            # Extract metadata from path and content
            extracted_metadata = self._extract_metadata_from_path(file_path)
            content_metadata = self._extract_metadata_from_content(text, file_ext)
            
            # Build document object for indexing
            document = {
                # Core fields
                'id': file_metadata.get('id'),
                'name': file_metadata.get('name'),
                'file_path': file_path,
                'content': text[:50000],  # Limit content size for main document
                'full_text': text,  # Keep full text for chunking (not stored in main doc)
                
                # Extracted metadata (dynamic, not hardcoded)
                'project_name': extracted_metadata.get('project'),
                'contractor': extracted_metadata.get('contractor'),
                'document_type': self._infer_document_type(file_path, text),
                
                # File metadata
                'file_size': file_metadata.get('size', 0),
                'modified_date': file_metadata.get('server_modified'),
                'created_date': file_metadata.get('client_modified'),
                
                # Content-based metadata
                'invoice_number': content_metadata.get('invoice_number'),
                'invoice_amount': content_metadata.get('amount'),
                'invoice_date': content_metadata.get('date'),
                'vendor_name': content_metadata.get('vendor'),
                
                # Processing metadata
                'indexed_at': datetime.utcnow().isoformat(),
                'text_length': len(text),
                'word_count': len(text.split())
            }
            
            # Clean None values
            document = {k: v for k, v in document.items() if v is not None}
            
            self.processed_count += 1
            return document
            
        except Exception as e:
            logger.error(f"Failed to process document {file_metadata.get('name')}: {e}")
            self.failed_count += 1
            return None
    
    def _extract_pdf_text(self, content: bytes) -> Optional[str]:
        """Extract text from PDF content"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = '\n'.join(text_parts)
            
            # Clean up PDF extraction issues while preserving structure
            # 1. Normalize line endings (CRLF -> LF)
            full_text = full_text.replace('\r\n', '\n').replace('\r', '\n')
            
            # 2. Fix hyphenation across lines FIRST (before other processing)
            full_text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', full_text)
            
            # 3. Normalize spaces/tabs within lines (but keep newlines!)
            full_text = re.sub(r'[ \t]+', ' ', full_text)
            
            # 4. Collapse excessive blank lines (keep max 2 newlines)
            full_text = re.sub(r'\n{3,}', '\n\n', full_text)
            
            # 5. Clean up spaces around newlines
            full_text = re.sub(r' *\n *', '\n', full_text)
            
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return None
    
    def _extract_docx_text(self, content: bytes) -> Optional[str]:
        """
        Extract text from DOCX file using python-docx
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Extracted text or None
        """
        if not docx_available:
            logger.warning("python-docx not available - cannot process DOCX")
            return None
            
        try:
            # Use python-docx to extract text from bytes (works reliably with BytesIO)
            doc = Document(io.BytesIO(content))
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            # Join all text
            full_text = '\n'.join(text_parts)
            
            # Clean up whitespace while preserving document structure
            # 1. Normalize spaces/tabs within lines (but keep newlines!)
            full_text = re.sub(r'[ \t]+', ' ', full_text)
            
            # 2. Collapse excessive blank lines (keep max 2 newlines)
            full_text = re.sub(r'\n{3,}', '\n\n', full_text)
            
            # 3. Clean up spaces around newlines
            full_text = re.sub(r' *\n *', '\n', full_text)
            
            full_text = full_text.strip()
            
            return full_text if full_text else None
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return None
    
    def _extract_xlsx_text(self, content: bytes) -> Optional[str]:
        """
        Extract text from XLSX file
        
        Args:
            content: XLSX file content as bytes
            
        Returns:
            Extracted text or None
        """
        if not openpyxl:
            logger.warning("openpyxl not available - cannot process XLSX")
            return None
            
        try:
            # Load workbook from bytes
            workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            
            text_parts = []
            
            # Extract text from all sheets
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}\n")
                
                # Extract cell values
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty cells and convert to string
                    row_values = [str(cell) for cell in row if cell is not None]
                    if row_values:
                        text_parts.append(' | '.join(row_values))
            
            workbook.close()
            
            # Join all text parts
            full_text = '\n'.join(text_parts)
            
            return full_text.strip() if full_text else None
            
        except Exception as e:
            logger.error(f"XLSX extraction failed: {e}")
            return None
    
    def _extract_metadata_from_path(self, path: str) -> Dict[str, Any]:
        """
        Extract metadata from file path
        Dynamic extraction - no hardcoded project names
        
        Args:
            path: File path
            
        Returns:
            Extracted metadata
        """
        metadata = {}
        parts = Path(path).parts
        
        # Look for patterns in path dynamically
        for i, part in enumerate(parts):
            # Projects often have numbers (addresses)
            if any(char.isdigit() for char in part) and i < 4:
                if 'project' not in metadata:
                    metadata['project'] = part
            
            # Look for contractor indicators
            if i > 0:
                prev_part = parts[i-1].upper()
                if any(indicator in prev_part for indicator in ['HIRED', 'CONTRACTOR', 'VENDOR']):
                    metadata['contractor'] = part
            
            # Document type indicators
            part_lower = part.lower()
            for doc_type in ['invoice', 'contract', 'agreement', 'report', 'w9', 'insurance']:
                if doc_type in part_lower:
                    metadata['document_type'] = doc_type
                    break
        
        return metadata
    
    def _extract_metadata_from_content(self, text: str, file_ext: str) -> Dict[str, Any]:
        """
        Extract metadata from document content
        Uses patterns to find invoice numbers, amounts, dates, etc.
        
        Args:
            text: Document text
            file_ext: File extension
            
        Returns:
            Extracted metadata
        """
        metadata = {}
        text_lower = text.lower()
        
        # Invoice number patterns (flexible, not vendor-specific)
        invoice_patterns = [
            r'invoice\s*#?\s*:?\s*([A-Z0-9-]+)',
            r'inv\s*#?\s*:?\s*([A-Z0-9-]+)',
            r'bill\s*#?\s*:?\s*([A-Z0-9-]+)',
            r'reference\s*:?\s*([A-Z0-9-]+)'
        ]
        
        for pattern in invoice_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metadata['invoice_number'] = match.group(1)
                break
        
        # Amount patterns (looking for dollar amounts)
        amount_patterns = [
            r'total\s*:?\s*\$?\s*([\d,]+\.?\d*)',
            r'amount\s*due\s*:?\s*\$?\s*([\d,]+\.?\d*)',
            r'balance\s*:?\s*\$?\s*([\d,]+\.?\d*)',
            r'\$\s*([\d,]+\.?\d*)'  # Any dollar amount
        ]
        
        for pattern in amount_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Take the largest amount (likely the total)
                amounts = []
                for match in matches:
                    try:
                        amount = float(match.replace(',', ''))
                        amounts.append(amount)
                    except:
                        pass
                if amounts:
                    metadata['amount'] = max(amounts)
                    break
        
        # Date patterns (flexible date extraction)
        date_patterns = [
            r'date\s*:?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'dated?\s*:?\s*(\w+\s+\d{1,2},?\s+\d{4})',
            r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                metadata['date'] = match.group(1)
                break
        
        # Vendor/Company name extraction - more flexible patterns
        # Pattern 1: Look for explicit vendor keywords (don't require newline at end)
        vendor_patterns = [
            r'(?:from|vendor|contractor|company|remit to|bill from|sold by|supplier|payee)\s*:?\s*([A-Za-z0-9\s&,.\-\']+?)(?:\n|\r|$|(?=[A-Z][a-z]*\s*:))',
            r'(?:from|vendor|contractor|company)\s*:?\s*([A-Za-z0-9\s&,.\-\']{3,100}?)[\n\r]',  # Keep original with newline as fallback
        ]
        
        vendor_found = False
        for pattern in vendor_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                vendor = match.group(1).strip()
                # Clean up vendor name
                vendor = re.sub(r'\s+', ' ', vendor)
                vendor = vendor.strip(' ,.-')  # Remove trailing punctuation
                if len(vendor) > 3 and len(vendor) < 100:  # Reasonable length
                    metadata['vendor'] = vendor
                    vendor_found = True
                    break
        
        # Pattern 2: Fallback heuristic - scan top lines for company-like text,
        # but avoid picking lines in the Bill To/Ship To block
        if not vendor_found:
            lines_full = [ln.strip() for ln in text.split('\n')]
            # Consider top N lines where header info typically appears
            N = min(40, len(lines_full))
            lines = lines_full[:N]

            # Identify and exclude the Bill To/Ship To block (next few lines after the marker)
            exclude_idx = set()
            for i, ln in enumerate(lines):
                if re.match(r'^(bill to|ship to)\b', ln, re.IGNORECASE):
                    for j in range(i+1, min(i+6, N)):
                        exclude_idx.add(j)

            # Build list of candidate indices with company indicators
            company_indicator = re.compile(r'\b(LLC|L\.L\.C\.|Inc|Inc\.|Incorporated|Corp|Corporation|Company|Co\.|Ltd|Limited|Partners|Partnership|Group|Associates|Enterprises)\b', re.IGNORECASE)

            candidates = []
            for idx, line in enumerate(lines):
                if not line or len(line) < 3:
                    continue
                # Skip explicit non-company headers and excluded Bill To/Ship To block
                if idx in exclude_idx or re.match(r'^(invoice|remit to|date|amount|total|subtotal|tax|terms|due on receipt|balance due)\b', line, re.IGNORECASE):
                    continue
                if re.match(r'^\d+[/-]\d+[/-]\d+', line) or re.match(r'^\$[\d,]+', line) or re.match(r'^page \d+', line, re.IGNORECASE) or re.match(r'^\d{3,}$', line):
                    continue
                if company_indicator.search(line):
                    candidates.append((idx, line))

            # Prefer candidates not in Bill To block and appearing near common vendor context (e.g., around Terms)
            def score_candidate(idx: int, line: str) -> int:
                score = 0
                # Prefer not excluded
                if idx not in exclude_idx:
                    score += 2
                # Prefer proximity to a 'Terms' line within next 5 lines (common on invoices)
                for k in range(idx, min(idx+6, N)):
                    if re.match(r'^terms\b', lines[k], re.IGNORECASE):
                        score += 2
                        break
                # Prefer lines following a 'Property Address' or address-like area
                for back in range(max(0, idx-5), idx):
                    if re.search(r'(address|suite|road|rd\.|street|st\.|texas|tx|zip)', lines[back], re.IGNORECASE):
                        score += 1
                        break
                return score

            if candidates:
                # Pick best-scoring candidate
                best = max(candidates, key=lambda t: score_candidate(t[0], t[1]))
                vendor = re.sub(r'\s+', ' ', best[1]).strip(' ,.-:')
                if 3 < len(vendor) < 100:
                    metadata['vendor'] = vendor
                    vendor_found = True

        return metadata
    
    def _infer_document_type(self, path: str, content: str) -> Optional[str]:
        """
        Infer document type from path and content
        Dynamic inference, not hardcoded
        
        Args:
            path: File path
            content: Document content
            
        Returns:
            Inferred document type
        """
        path_lower = path.lower()
        content_lower = content.lower()[:1000]  # Check first 1000 chars
        
        # Check for document type indicators
        type_indicators = {
            'invoice': ['invoice', 'bill', 'statement', 'remittance'],
            'contract': ['agreement', 'contract', 'terms and conditions', 'contractor'],
            'report': ['report', 'analysis', 'assessment', 'evaluation', 'summary'],
            'w9': ['w-9', 'w9', 'taxpayer identification', 'tin', 'form w'],
            'insurance': ['insurance', 'certificate of insurance', 'coi', 'liability', 'coverage'],
            'receipt': ['receipt', 'payment received', 'paid', 'transaction'],
            'change_order': ['change order', 'modification', 'amendment', 'variation']
        }
        
        for doc_type, indicators in type_indicators.items():
            for indicator in indicators:
                if indicator in path_lower or indicator in content_lower:
                    return doc_type
        
        return None
    
    def chunk_text(self, text: str, chunk_size: int = 1000, 
                   overlap: int = 200) -> List[str]:
        """
        Chunk text for vector embedding
        Creates overlapping chunks for better context
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk in characters
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + chunk_size, text_length)
            
            # Try to end at a sentence boundary
            if end < text_length:
                # Look for sentence end
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap if end < text_length else text_length
        
        return chunks
    
    def get_processing_stats(self) -> Dict[str, int]:
        """Get statistics about document processing"""
        return {
            'processed': self.processed_count,
            'failed': self.failed_count,
            'total': self.processed_count + self.failed_count
        }
